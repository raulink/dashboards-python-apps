import dash
#from dash.dash_table.Format import Group
from dash import dash_table
#import dash_table
#import dash_core_components as dcc
#import dash_html_components as html
from dash import html, dcc
import dash_bootstrap_components as dbc
from dash.dependencies import Input, Output, State
import pandas as pd
from io import BytesIO
import base64
from docx import Document

# Inicializar la app de Dash con Bootstrap
app = dash.Dash(__name__, external_stylesheets=[dbc.themes.BOOTSTRAP])

app.layout = html.Div([
    html.H1("Sistema de Contrataciones"),
    
    # Botón para subir el archivo Excel
    dcc.Upload(
        id='upload-excel',
        children=html.Button('Subir Archivo Excel'),
        style={'width': '100%', 'height': '60px', 'lineHeight': '60px', 'borderWidth': '1px', 'borderStyle': 'dashed', 'borderRadius': '5px', 'textAlign': 'center', 'margin': '10px'}
    ),
    
    # Tabla para visualizar y editar los datos del Excel
    dash_table.DataTable(
        id='table',
        editable=True,
    ),
    
    # Botón para subir la plantilla de Word
    dcc.Upload(
        id='upload-template',
        children=html.Button('Subir Plantilla', disabled=True),
        style={'width': '100%', 'height': '60px', 'lineHeight': '60px', 'borderWidth': '1px', 'borderStyle': 'dashed', 'borderRadius': '5px', 'textAlign': 'center', 'margin': '10px'}
    ),
    
    # Botón para descargar los archivos generados
    html.Button("Descargar Proceso", id="download-button", n_clicks=0, disabled=True),
    dcc.Download(id="download-excel"),
    dcc.Download(id="download-word"),
])

@app.callback(
    [Output('table', 'data'), Output('table', 'columns'), Output('upload-template', 'children')],
    [Input('upload-excel', 'contents')],
    [State('upload-excel', 'filename')]
)
def update_table(contents, filename):
    if contents is None:
        return [], [], html.Button('Subir Plantilla', disabled=True)
    
    content_type, content_string = contents.split(',')
    decoded = base64.b64decode(content_string)
    df = pd.read_excel(BytesIO(decoded), sheet_name='Hoja1')
    
    data = df.to_dict('records')
    columns = [{'name': col, 'id': col} for col in df.columns]
    
    return data, columns, html.Button('Subir Plantilla')

@app.callback(
    Output('download-button', 'disabled'),
    [Input('upload-template', 'contents')]
)
def enable_download_button(contents):
    if contents is not None:
        return False
    return True

@app.callback(
    [Output('download-excel', 'data'), Output('download-word', 'data')],
    [Input('download-button', 'n_clicks')],
    [State('table', 'data'), State('upload-template', 'contents')]
)
def generate_files(n_clicks, table_data, template_contents):
    if n_clicks == 0:
        return None, None
    
    df = pd.DataFrame(table_data)
    
    # Generar archivo Excel
    excel_buffer = BytesIO()
    df.to_excel(excel_buffer, index=False)
    excel_data = dcc.send_data_frame(df.to_excel, "output.xlsx", sheet_name="Hoja1", index=False)
    
    # Generar archivo Word
    content_type, content_string = template_contents.split(',')
    decoded = base64.b64decode(content_string)
    template = Document(BytesIO(decoded))
    
    # Aquí puedes modificar el documento de Word según necesites
    template.add_paragraph("Datos de la tabla:")
    for index, row in df.iterrows():
        template.add_paragraph(str(row.to_dict()))
    
    word_buffer = BytesIO()
    template.save(word_buffer)
    word_data = dict(content=word_buffer.getvalue(), filename="output.docx")
    
    return excel_data, word_data

if __name__ == '__main__':
    app.run_server(debug=True)
